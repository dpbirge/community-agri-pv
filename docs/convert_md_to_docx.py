"""Convert markdown files to Word documents with formatting."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def parse_markdown(md_text):
    """Parse markdown into a list of block elements."""
    blocks = []
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            blocks.append(('code', '\n'.join(code_lines)))
            i += 1
            continue

        # Heading
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append(('heading', level, heading_match.group(2).strip()))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+\s*$', line):
            blocks.append(('hr',))
            i += 1
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append(('table', table_lines))
            continue

        # List item
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if list_match:
            indent = len(list_match.group(1))
            level = indent // 2
            is_ordered = bool(re.match(r'\d+\.', list_match.group(2)))
            blocks.append(('list_item', level, is_ordered, list_match.group(3).strip()))
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        blocks.append(('paragraph', line.strip()))
        i += 1

    return blocks


def add_formatted_text(paragraph, text):
    """Add text with inline markdown formatting (bold, italic, code)."""
    # Split on inline formatting patterns
    pattern = r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_table_row(row_text):
    """Parse a markdown table row into cells."""
    cells = row_text.strip().strip('|').split('|')
    return [c.strip() for c in cells]


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a Word document."""
    md_text = Path(md_path).read_text()
    blocks = parse_markdown(md_text)
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for block in blocks:
        btype = block[0]

        if btype == 'heading':
            level = block[1]
            text = block[2]
            heading = doc.add_heading(level=min(level, 4))
            add_formatted_text(heading, text)

        elif btype == 'paragraph':
            p = doc.add_paragraph()
            add_formatted_text(p, block[1])

        elif btype == 'list_item':
            level, is_ordered, text = block[1], block[2], block[3]
            style_name = 'List Number' if is_ordered else 'List Bullet'
            p = doc.add_paragraph(style=style_name)
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            add_formatted_text(p, text)

        elif btype == 'code':
            code_text = block[1]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elif btype == 'table':
            table_lines = block[1]
            # Skip separator row
            header_cells = parse_table_row(table_lines[0])
            data_rows = [parse_table_row(line) for line in table_lines[2:]]
            num_cols = len(header_cells)

            table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
            table.style = 'Table Grid'

            # Header row
            for j, cell_text in enumerate(header_cells):
                cell = table.rows[0].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(10)
                set_cell_shading(cell, 'D9E2F3')

            # Data rows
            for i, row_data in enumerate(data_rows):
                for j, cell_text in enumerate(row_data):
                    if j < num_cols:
                        cell = table.rows[i + 1].cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        add_formatted_text(p, cell_text)
                        for run in p.runs:
                            run.font.size = Pt(10)

            doc.add_paragraph()  # spacing after table

        elif btype == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin horizontal line via border
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): 'auto',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    base = Path(__file__).parent.parent
    out_dir = Path(__file__).parent

    convert_md_to_docx(base / 'structure.md', out_dir / 'structure.docx')
    convert_md_to_docx(base / 'policies.md', out_dir / 'policies.docx')
